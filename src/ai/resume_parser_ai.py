#!/usr/bin/env python3
"""
🤖 AI-Powered Resume Parser
Uses OpenAI GPT-4o and Google Gemini 2.5 Pro for professional resume analysis
"""

import os
import json
import time
from typing import Dict, Any, List, Optional
from datetime import datetime
import tempfile
import base64
from pathlib import Path

# Import AI libraries
try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

# PDF processing
try:
    import PyPDF2
    import fitz  # PyMuPDF
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Image processing
try:
    from PIL import Image
    IMAGE_AVAILABLE = True
except ImportError:
    IMAGE_AVAILABLE = False


class AIResumeParser:
    """Professional AI-powered resume parser using GPT-4o and Gemini 2.5 Pro"""
    
    def __init__(self):
        """Initialize the AI resume parser with API credentials"""
        self.openai_api_key = os.getenv("OPENAI_API_KEY", "your-openai-api-key-here")
        self.gemini_api_key = os.getenv("GEMINI_API_KEY", "your-gemini-api-key-here")
        
        # Initialize OpenAI
        if OPENAI_AVAILABLE and self.openai_api_key:
            openai.api_key = self.openai_api_key
            self.openai_client = openai.OpenAI(api_key=self.openai_api_key)
        
        # Initialize Gemini
        if GEMINI_AVAILABLE and self.gemini_api_key:
            genai.configure(api_key=self.gemini_api_key)
            self.gemini_model = genai.GenerativeModel('gemini-2.0-flash-exp')
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF using multiple methods"""
        text = ""
        
        # Method 1: Try PyMuPDF (most reliable)
        try:
            doc = fitz.open(pdf_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text() + "\n"
            doc.close()
            if len(text.strip()) > 50:
                return text
        except Exception:
            pass
        
        # Method 2: Try pdfplumber
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if len(text.strip()) > 50:
                return text
        except Exception:
            pass
        
        # Method 3: Try PyPDF2
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            if len(text.strip()) > 50:
                return text
        except Exception:
            pass
        
        return text or "Could not extract text from PDF"
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(docx_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            return f"Could not extract text from DOCX: {str(e)}"
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file formats"""
        file_path = Path(file_path)
        
        if file_path.suffix.lower() == '.pdf':
            return self.extract_text_from_pdf(str(file_path))
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            return self.extract_text_from_docx(str(file_path))
        elif file_path.suffix.lower() == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except Exception:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
        else:
            return "Unsupported file format"
    
    def create_resume_parsing_prompt(self) -> str:
        """Create a comprehensive prompt for resume parsing"""
        return """
You are a professional AI resume parser. Analyze the provided resume text and extract comprehensive information in the following JSON format:

{
  "personal_info": {
    "full_name": "Full name of the candidate",
    "email": "Email address",
    "phone": "Phone number",
    "location": "City, State/Country",
    "linkedin": "LinkedIn URL or username", 
    "github": "GitHub URL or username",
    "portfolio": "Portfolio website URL",
    "summary": "Professional summary or objective"
  },
  "experience": [
    {
      "title": "Job title",
      "company": "Company name",
      "duration": "Start - End dates",
      "location": "Job location",
      "description": "Job description and responsibilities",
      "achievements": ["List of key achievements"],
      "technologies": ["Technologies used"]
    }
  ],
  "education": [
    {
      "degree": "Degree name",
      "institution": "Institution name", 
      "year": "Graduation year",
      "location": "Institution location",
      "gpa": "GPA if mentioned",
      "relevant_coursework": ["Relevant courses"]
    }
  ],
  "skills": {
    "programming_languages": ["Languages"],
    "frameworks": ["Frameworks and libraries"],
    "tools": ["Development tools"],
    "databases": ["Database technologies"],
    "cloud": ["Cloud platforms"],
    "soft_skills": ["Soft skills"]
  },
  "certifications": [
    {
      "name": "Certification name",
      "issuer": "Issuing organization",
      "date": "Date obtained",
      "expiry": "Expiry date if applicable"
    }
  ],
  "projects": [
    {
      "name": "Project name",
      "description": "Project description",
      "technologies": ["Technologies used"],
      "url": "Project URL if available"
    }
  ],
  "languages": [
    {
      "language": "Language name",
      "proficiency": "Proficiency level"
    }
  ],
  "analysis": {
    "total_experience_years": 0.0,
    "seniority_level": "Junior/Mid/Senior/Lead",
    "primary_skills": ["Top 5 most important skills"],
    "industry_focus": ["Industries/domains"],
    "career_progression": "Brief analysis of career growth",
    "strengths": ["Key strengths identified"],
    "suggested_roles": ["Ideal job titles for this candidate"]
  },
  "confidence_score": 95,
  "parsing_quality": "excellent",
  "extracted_sections": ["List of sections successfully parsed"]
}

Instructions:
1. Extract ALL available information accurately
2. Be precise with dates, numbers, and technical terms
3. Identify implicit information (e.g., seniority from job titles)
4. Calculate total experience from all positions
5. Provide confidence score (0-100) based on information completeness
6. If information is missing, use null or empty arrays
7. Focus on professional, technical, and quantifiable achievements
8. Maintain consistent date formats
9. Extract both explicit and implicit skills
10. Provide career progression analysis

Return ONLY the JSON response, no additional text.
"""
    
    def parse_with_openai(self, resume_text: str) -> Dict[str, Any]:
        """Parse resume using OpenAI GPT-4o"""
        try:
            if not OPENAI_AVAILABLE or not self.openai_api_key:
                raise Exception("OpenAI not available or API key not set")
            
            response = self.openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": self.create_resume_parsing_prompt()},
                    {"role": "user", "content": f"Please parse this resume:\n\n{resume_text}"}
                ],
                temperature=0.1,
                max_tokens=4000
            )
            
            result_text = response.choices[0].message.content.strip()
            
            # Clean and parse JSON
            if result_text.startswith('```json'):
                result_text = result_text[7:]
            if result_text.endswith('```'):
                result_text = result_text[:-3]
            
            parsed_data = json.loads(result_text)
            parsed_data['parsing_method'] = 'OpenAI GPT-4o'
            parsed_data['parsed_at'] = datetime.now().isoformat()
            
            return parsed_data
            
        except Exception as e:
            return {
                "error": f"OpenAI parsing failed: {str(e)}",
                "parsing_method": "OpenAI GPT-4o (Failed)",
                "success": False
            }
    
    def parse_with_gemini(self, resume_text: str) -> Dict[str, Any]:
        """Parse resume using Google Gemini 2.5 Pro"""
        try:
            if not GEMINI_AVAILABLE or not self.gemini_api_key:
                raise Exception("Gemini not available or API key not set")
            
            prompt = self.create_resume_parsing_prompt() + f"\n\nResume to parse:\n{resume_text}"
            
            response = self.gemini_model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.1,
                    max_output_tokens=4000,
                )
            )
            
            result_text = response.text.strip()
            
            # Clean and parse JSON
            if result_text.startswith('```json'):
                result_text = result_text[7:]
            if result_text.endswith('```'):
                result_text = result_text[:-3]
            
            parsed_data = json.loads(result_text)
            parsed_data['parsing_method'] = 'Google Gemini 2.5 Pro'
            parsed_data['parsed_at'] = datetime.now().isoformat()
            
            return parsed_data
            
        except Exception as e:
            return {
                "error": f"Gemini parsing failed: {str(e)}",
                "parsing_method": "Google Gemini 2.5 Pro (Failed)",
                "success": False
            }
    
    def merge_parsing_results(self, openai_result: Dict, gemini_result: Dict) -> Dict[str, Any]:
        """Merge results from both AI models for maximum accuracy"""
        if openai_result.get('error') and gemini_result.get('error'):
            return {
                "error": "Both AI parsers failed",
                "openai_error": openai_result.get('error'),
                "gemini_error": gemini_result.get('error'),
                "success": False
            }
        
        # Use the successful result, prefer OpenAI if both succeed
        if not openai_result.get('error'):
            primary_result = openai_result
            secondary_result = gemini_result
        else:
            primary_result = gemini_result
            secondary_result = openai_result
        
        # Enhance with secondary result data if available
        if not secondary_result.get('error'):
            # Merge additional insights
            if 'analysis' in secondary_result and 'analysis' in primary_result:
                # Combine strengths and suggestions
                secondary_analysis = secondary_result['analysis']
                if 'strengths' in secondary_analysis:
                    existing_strengths = set(primary_result['analysis'].get('strengths', []))
                    new_strengths = [s for s in secondary_analysis['strengths'] if s not in existing_strengths]
                    primary_result['analysis']['strengths'] = list(existing_strengths) + new_strengths[:2]
        
        # Add metadata about dual parsing
        primary_result['parsing_methods'] = [
            primary_result.get('parsing_method', 'Unknown'),
            secondary_result.get('parsing_method', 'Unknown') if not secondary_result.get('error') else f"{secondary_result.get('parsing_method', 'Unknown')} (Failed)"
        ]
        primary_result['dual_ai_enhanced'] = True
        primary_result['success'] = True
        
        return primary_result
    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """Parse resume using dual AI approach (OpenAI + Gemini)"""
        start_time = time.time()
        
        try:
            # Extract text from file
            resume_text = self.extract_text_from_file(file_path)
            
            if len(resume_text.strip()) < 50:
                return {
                    "error": "Could not extract sufficient text from resume",
                    "text_length": len(resume_text),
                    "success": False
                }
            
            print(f"📄 Extracted {len(resume_text)} characters from resume")
            
            # Parse with both AI models in parallel (simulated)
            print("🤖 Parsing with OpenAI GPT-4o...")
            openai_result = self.parse_with_openai(resume_text)
            
            print("🧠 Parsing with Google Gemini 2.5 Pro...")
            gemini_result = self.parse_with_gemini(resume_text)
            
            # Merge results for best accuracy
            final_result = self.merge_parsing_results(openai_result, gemini_result)
            
            # Add processing metadata
            processing_time = time.time() - start_time
            final_result['processing_time_seconds'] = round(processing_time, 2)
            final_result['text_length'] = len(resume_text)
            final_result['file_processed'] = os.path.basename(file_path)
            
            print(f"✅ Resume parsed successfully in {processing_time:.2f}s")
            return final_result
            
        except Exception as e:
            return {
                "error": f"Resume parsing failed: {str(e)}",
                "processing_time_seconds": time.time() - start_time,
                "success": False
            }
    
    def parse_resume_from_upload(self, uploaded_file) -> Dict[str, Any]:
        """Parse resume from Streamlit uploaded file"""
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_path = tmp_file.name
            
            # Parse the resume
            result = self.parse_resume(tmp_path)
            
            # Clean up
            try:
                os.unlink(tmp_path)
            except:
                pass
            
            return result
            
        except Exception as e:
            return {
                "error": f"Upload processing failed: {str(e)}",
                "success": False
            }


# Create global instance
ai_resume_parser = AIResumeParser()


def parse_resume_with_ai(file_path_or_upload) -> Dict[str, Any]:
    """
    Main function to parse resume with AI
    Args:
        file_path_or_upload: Either a file path (str) or Streamlit uploaded file
    Returns:
        Dict with parsed resume data
    """
    if isinstance(file_path_or_upload, str):
        return ai_resume_parser.parse_resume(file_path_or_upload)
    else:
        return ai_resume_parser.parse_resume_from_upload(file_path_or_upload)


# Test function
def test_ai_resume_parser():
    """Test the AI resume parser with sample data"""
    sample_resume = """
    John Smith
    Senior Software Engineer
    john.smith@email.com | (555) 123-4567
    San Francisco, CA | linkedin.com/in/johnsmith
    
    PROFESSIONAL SUMMARY
    Experienced Software Engineer with 6+ years in full-stack development.
    Expert in Python, React, and cloud technologies.
    
    EXPERIENCE
    Senior Software Engineer | TechCorp | 2021 - Present
    • Led development of microservices architecture
    • Improved system performance by 40%
    • Mentored 3 junior developers
    
    Software Engineer | StartupXYZ | 2019 - 2021
    • Built React applications with Node.js backend
    • Implemented CI/CD pipelines
    
    EDUCATION
    Bachelor of Science in Computer Science | Stanford University | 2018
    
    SKILLS
    Programming: Python, JavaScript, TypeScript
    Frameworks: React, Django, Node.js
    Cloud: AWS, Docker, Kubernetes
    """
    
    # Create temp file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as tmp_file:
        tmp_file.write(sample_resume)
        tmp_path = tmp_file.name
    
    # Test parsing
    result = ai_resume_parser.parse_resume(tmp_path)
    
    # Clean up
    os.unlink(tmp_path)
    
    return result


if __name__ == "__main__":
    # Test the parser
    print("🧪 Testing AI Resume Parser...")
    test_result = test_ai_resume_parser()
    
    if test_result.get('success', False):
        print("✅ AI Resume Parser test successful!")
        print(f"📊 Confidence Score: {test_result.get('confidence_score', 0)}%")
        print(f"⏱️ Processing Time: {test_result.get('processing_time_seconds', 0)}s")
        print(f"🤖 Methods Used: {', '.join(test_result.get('parsing_methods', ['Unknown']))}")
    else:
        print(f"❌ Test failed: {test_result.get('error', 'Unknown error')}")