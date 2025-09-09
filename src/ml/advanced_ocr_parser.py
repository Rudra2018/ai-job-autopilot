#!/usr/bin/env python3
"""
🔍 Advanced OCR Resume Parser
Multi-engine OCR with Google Vision, Tesseract, and EasyOCR for maximum accuracy
"""

import logging
import tempfile
import os
from typing import Dict, List, Any, Optional, Tuple
import json
import re
from pathlib import Path

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# OCR Engine Availability
GOOGLE_VISION_AVAILABLE = False
TESSERACT_AVAILABLE = False
EASYOCR_AVAILABLE = False
PADDLEOCR_AVAILABLE = False

# Initialize OCR engines
try:
    from google.cloud import vision
    GOOGLE_VISION_AVAILABLE = True
    logger.info("✅ Google Vision OCR available")
except ImportError:
    logger.warning("⚠️ Google Vision not available - install with: pip install google-cloud-vision")

try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
    logger.info("✅ Tesseract OCR available")
except ImportError:
    logger.warning("⚠️ Tesseract not available - install with: pip install pytesseract pillow")

try:
    import easyocr
    EASYOCR_AVAILABLE = True
    logger.info("✅ EasyOCR available")
except ImportError:
    logger.warning("⚠️ EasyOCR not available - install with: pip install easyocr")

try:
    from paddleocr import PaddleOCR
    PADDLEOCR_AVAILABLE = True
    logger.info("✅ PaddleOCR available")
except ImportError:
    logger.warning("⚠️ PaddleOCR not available - install with: pip install paddlepaddle paddleocr")

class AdvancedOCRParser:
    """Advanced multi-engine OCR parser for maximum text extraction accuracy"""
    
    def __init__(self):
        self.ocr_engines = []
        self._initialize_engines()
    
    def _initialize_engines(self):
        """Initialize available OCR engines"""
        if GOOGLE_VISION_AVAILABLE:
            try:
                self.vision_client = vision.ImageAnnotatorClient()
                self.ocr_engines.append('google_vision')
                logger.info("🔍 Google Vision OCR initialized")
            except Exception as e:
                logger.warning(f"⚠️ Google Vision init failed: {e}")
        
        if TESSERACT_AVAILABLE:
            self.ocr_engines.append('tesseract')
        
        if EASYOCR_AVAILABLE:
            try:
                self.easyocr_reader = easyocr.Reader(['en'])
                self.ocr_engines.append('easyocr')
            except Exception as e:
                logger.warning(f"⚠️ EasyOCR init failed: {e}")
        
        if PADDLEOCR_AVAILABLE:
            try:
                self.paddle_ocr = PaddleOCR(use_angle_cls=True, lang='en')
                self.ocr_engines.append('paddleocr')
            except Exception as e:
                logger.warning(f"⚠️ PaddleOCR init failed: {e}")
        
        logger.info(f"🚀 Initialized {len(self.ocr_engines)} OCR engines: {self.ocr_engines}")
    
    def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """Extract text using multiple OCR engines and return best result"""
        results = {}
        
        # Determine file type
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return self._extract_from_image(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._extract_from_docx(file_path)
        elif file_ext == '.txt':
            return self._extract_from_txt(file_path)
        else:
            logger.warning(f"⚠️ Unsupported file type: {file_ext}")
            return {'text': '', 'confidence': 0, 'method': 'unsupported'}
    
    def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF using multiple methods"""
        results = []
        
        # Method 1: PyPDF2 for text-based PDFs
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                if len(text.strip()) > 100:  # Good text extraction
                    results.append({
                        'text': text,
                        'confidence': 95,
                        'method': 'pypdf2_text'
                    })
        except Exception as e:
            logger.warning(f"⚠️ PyPDF2 extraction failed: {e}")
        
        # Method 2: Convert PDF to images and use OCR
        try:
            import pdf2image
            from PIL import Image
            
            # Convert PDF to images
            pages = pdf2image.convert_from_path(file_path, dpi=300)
            
            all_text = ""
            total_confidence = 0
            
            for i, page in enumerate(pages[:5]):  # Process first 5 pages
                # Save page as temp image
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
                    page.save(tmp_img.name, 'PNG')
                    
                    # Extract text from image
                    img_result = self._extract_from_image(tmp_img.name)
                    all_text += img_result['text'] + "\n"
                    total_confidence += img_result['confidence']
                    
                    # Clean up
                    os.unlink(tmp_img.name)
            
            if pages:
                avg_confidence = total_confidence / len(pages[:5])
                results.append({
                    'text': all_text,
                    'confidence': avg_confidence,
                    'method': 'pdf_to_image_ocr'
                })
        
        except Exception as e:
            logger.warning(f"⚠️ PDF to image OCR failed: {e}")
        
        # Return best result
        if results:
            best_result = max(results, key=lambda x: x['confidence'])
            logger.info(f"📄 PDF extraction: {best_result['method']} (confidence: {best_result['confidence']}%)")
            return best_result
        else:
            return {'text': '', 'confidence': 0, 'method': 'pdf_failed'}
    
    def _extract_from_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text from image using multiple OCR engines"""
        results = []
        
        # Google Vision OCR
        if 'google_vision' in self.ocr_engines:
            try:
                result = self._google_vision_ocr(file_path)
                if result['confidence'] > 0:
                    results.append(result)
            except Exception as e:
                logger.warning(f"⚠️ Google Vision OCR failed: {e}")
        
        # EasyOCR
        if 'easyocr' in self.ocr_engines:
            try:
                result = self._easyocr_extract(file_path)
                if result['confidence'] > 0:
                    results.append(result)
            except Exception as e:
                logger.warning(f"⚠️ EasyOCR failed: {e}")
        
        # Tesseract OCR
        if 'tesseract' in self.ocr_engines:
            try:
                result = self._tesseract_extract(file_path)
                if result['confidence'] > 0:
                    results.append(result)
            except Exception as e:
                logger.warning(f"⚠️ Tesseract OCR failed: {e}")
        
        # PaddleOCR
        if 'paddleocr' in self.ocr_engines:
            try:
                result = self._paddleocr_extract(file_path)
                if result['confidence'] > 0:
                    results.append(result)
            except Exception as e:
                logger.warning(f"⚠️ PaddleOCR failed: {e}")
        
        # Return best result
        if results:
            best_result = max(results, key=lambda x: x['confidence'])
            logger.info(f"🖼️ Image extraction: {best_result['method']} (confidence: {best_result['confidence']}%)")
            return best_result
        else:
            return {'text': '', 'confidence': 0, 'method': 'image_failed'}
    
    def _google_vision_ocr(self, file_path: str) -> Dict[str, Any]:
        """Extract text using Google Vision API"""
        try:
            with open(file_path, 'rb') as image_file:
                content = image_file.read()
            
            image = vision.Image(content=content)
            response = self.vision_client.text_detection(image=image)
            
            if response.error.message:
                raise Exception(response.error.message)
            
            texts = response.text_annotations
            if texts:
                extracted_text = texts[0].description
                # Calculate confidence based on text quality
                confidence = self._calculate_text_confidence(extracted_text)
                
                return {
                    'text': extracted_text,
                    'confidence': min(95, confidence + 10),  # Google Vision bonus
                    'method': 'google_vision'
                }
        
        except Exception as e:
            logger.error(f"❌ Google Vision OCR error: {e}")
        
        return {'text': '', 'confidence': 0, 'method': 'google_vision_failed'}
    
    def _easyocr_extract(self, file_path: str) -> Dict[str, Any]:
        """Extract text using EasyOCR"""
        try:
            results = self.easyocr_reader.readtext(file_path, detail=1)
            
            text_parts = []
            total_confidence = 0
            
            for (bbox, text, conf) in results:
                if conf > 0.3:  # Filter low confidence
                    text_parts.append(text)
                    total_confidence += conf
            
            extracted_text = ' '.join(text_parts)
            avg_confidence = (total_confidence / len(results)) * 100 if results else 0
            
            return {
                'text': extracted_text,
                'confidence': int(avg_confidence),
                'method': 'easyocr'
            }
        
        except Exception as e:
            logger.error(f"❌ EasyOCR error: {e}")
        
        return {'text': '', 'confidence': 0, 'method': 'easyocr_failed'}
    
    def _tesseract_extract(self, file_path: str) -> Dict[str, Any]:
        """Extract text using Tesseract OCR"""
        try:
            from PIL import Image
            import pytesseract
            
            image = Image.open(file_path)
            
            # Get text and confidence data
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            
            # Filter confident text
            confident_text = []
            total_conf = 0
            valid_words = 0
            
            for i, conf in enumerate(data['conf']):
                if int(conf) > 30:  # Filter low confidence
                    word = data['text'][i].strip()
                    if word:
                        confident_text.append(word)
                        total_conf += int(conf)
                        valid_words += 1
            
            extracted_text = ' '.join(confident_text)
            avg_confidence = total_conf / valid_words if valid_words > 0 else 0
            
            return {
                'text': extracted_text,
                'confidence': int(avg_confidence),
                'method': 'tesseract'
            }
        
        except Exception as e:
            logger.error(f"❌ Tesseract error: {e}")
        
        return {'text': '', 'confidence': 0, 'method': 'tesseract_failed'}
    
    def _paddleocr_extract(self, file_path: str) -> Dict[str, Any]:
        """Extract text using PaddleOCR"""
        try:
            result = self.paddle_ocr.ocr(file_path, cls=True)
            
            text_parts = []
            total_confidence = 0
            valid_parts = 0
            
            for line in result:
                for word_info in line:
                    text = word_info[1][0]
                    confidence = word_info[1][1]
                    
                    if confidence > 0.5:  # Filter low confidence
                        text_parts.append(text)
                        total_confidence += confidence
                        valid_parts += 1
            
            extracted_text = ' '.join(text_parts)
            avg_confidence = (total_confidence / valid_parts) * 100 if valid_parts > 0 else 0
            
            return {
                'text': extracted_text,
                'confidence': int(avg_confidence),
                'method': 'paddleocr'
            }
        
        except Exception as e:
            logger.error(f"❌ PaddleOCR error: {e}")
        
        return {'text': '', 'confidence': 0, 'method': 'paddleocr_failed'}
    
    def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return {
                'text': text,
                'confidence': 100,  # DOCX text extraction is reliable
                'method': 'docx_native'
            }
        
        except Exception as e:
            logger.error(f"❌ DOCX extraction error: {e}")
        
        return {'text': '', 'confidence': 0, 'method': 'docx_failed'}
    
    def _extract_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                'text': text,
                'confidence': 100,  # Plain text is reliable
                'method': 'txt_native'
            }
        
        except Exception as e:
            logger.error(f"❌ TXT extraction error: {e}")
        
        return {'text': '', 'confidence': 0, 'method': 'txt_failed'}
    
    def _calculate_text_confidence(self, text: str) -> int:
        """Calculate confidence score based on text quality metrics"""
        if not text or len(text.strip()) < 10:
            return 0
        
        score = 50  # Base score
        
        # Length bonus
        if len(text) > 100:
            score += 10
        if len(text) > 500:
            score += 10
        
        # Word count bonus
        words = text.split()
        if len(words) > 20:
            score += 10
        
        # Email pattern bonus
        if re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text):
            score += 5
        
        # Phone pattern bonus
        if re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', text):
            score += 5
        
        # Professional keywords bonus
        professional_keywords = [
            'experience', 'skills', 'education', 'work', 'employment',
            'project', 'achievement', 'responsibility', 'manage', 'develop'
        ]
        keyword_count = sum(1 for keyword in professional_keywords if keyword.lower() in text.lower())
        score += min(10, keyword_count * 2)
        
        return min(100, score)

# Global instance
ocr_parser = AdvancedOCRParser()

def extract_text_with_advanced_ocr(file_path: str) -> Dict[str, Any]:
    """Main function to extract text using advanced OCR"""
    return ocr_parser.extract_text_from_file(file_path)